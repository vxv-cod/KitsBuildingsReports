import os, sys
from requests import get as requests_get
from requests_negotiate_sspi import HttpNegotiateAuth

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT
# from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# from rich import print


# def decorTime(my_func):
#     '''Обертка функции декоратором'''
#     from time import time
#     # *args — это сокращение от «arguments» (аргументы) в виде кортежа, 
#     # **kwargs — сокращение от «keyword arguments» (именованные аргументы)
#     def wrapper(*args):
#         start_time = time()
#         my_func(*args)
#         print(f'"Выполнено: {round(time() - start_time, 2)} sec"')
#     return wrapper


def resp(url):
    response = requests_get(url, auth=HttpNegotiateAuth())
    if response.status_code == 404:
        # print("'Страница не существует!'")
        response = 'HELP!!!'
        return response
    if response.status_code == 401:
        print("'Ошибка авторизации пользователя!'")
        # response = None
        return 
    # print(response)
    response = response.json()
    return response


def get_dictionary_item_Title(Id):
    '''Список Id словарей'''
    global dictionary_items
    for i in dictionary_items:
        if i['Id'] == Id:
            return (i['Title'])


# def copy_elem(paragraph, elem):
#     from copy import deepcopy
#     paragraph._p.addnext(deepcopy(elem._element))


def delete_paragraph(paragraph):
    '''Удаление параграфа и таблицы'''
    p = paragraph._element
    p.getparent().remove(p)


def set_cell_color(cell):
    ''' Фон ячейки'''
    clShading = OxmlElement('w:shd')
    clShading.set(qn('w:fill'), "FFD200") # Hex of RGBColor(255, 210, 0)
    cell._tc.get_or_add_tcPr().append(clShading)
    return cell


def min_max(table, StartRow, Col):
    '''Инедксы крайних объединяемых ячеек по столбцам'''
    celMergeCol = table.column_cells(Col)
    celMergeCol = celMergeCol[StartRow:]
    celMergeCol = [i.text for i in celMergeCol]
    # print(celMergeCol)

    ind_min_max_list = []
    xxx = []
    for i in range(len(celMergeCol)):
        if celMergeCol[i] != 'None':
            if len(xxx) > 1:
                ind_min_max_list.append([min(xxx), max(xxx)])
                xxx = []
        if celMergeCol[i] == 'None':
            pass
        xxx.append(i + StartRow)
        if i == len(celMergeCol) - 1:
            ind_min_max_list.append([min(xxx), max(xxx)])
    return ind_min_max_list


def MegeCellColumns(table, ColMergeList, ind_min_max_list):
    '''Объединение ячеек в колонках таблице'''
    for Col in ColMergeList:
        for index in ind_min_max_list:
            Cell = table.cell(index[0], Col)
            other_cell = table.cell(index[1], Col)
            cell_merge = Cell.merge(other_cell)
            # Перезаписываем значение ячейки, оствляя только первый параграф в ней
            cell_merge.text = cell_merge.paragraphs[0].text


def MegeCellRows(table, row, col1, col2):
    Cell = table.cell(row, col1)
    other_cell = table.cell(row, col2)
    cell_merge = Cell.merge(other_cell)
    cell_merge.text = cell_merge.paragraphs[0].text
    return cell_merge.paragraphs[0]


def ExportdataInDoc(table, data):
    '''Присваиваем стиль таблицы'''
    table.style = 'vxv'    

    '''Добавляем данные в таблицу'''    
    for irow, row in enumerate(data):
        # добавляем строку с ячейками к объекту таблицы
        if irow == 0:
            cells = table.rows[-1]
        else:
            # добавляем строчку в таблицу
            cells = table.add_row()
        # Высота строки должна быть не менее минимального заданного значения
        cells.height_rule = WD_ROW_HEIGHT.AT_LEAST
        # Высота строки
        cells.height = Cm(0.8)
        # последовательность экземпляров Cell (к ячейкам строки)
        cells = cells.cells
        
        # вставляем данные в ячейки
        for i, item in enumerate(row):
            cells[i].text = str(item)
            # p = cells[i]
            # p.text = str(item)
            # p.paragraphs[0].font.color.rgb = RGBColor(255, 0, 0)           

            # Текст выровнен по центру ячейки
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            '''Выравнивание текста в ячейках через параграф'''
            # Подключаемся к параграфу в ячейке
            p = cells[i].paragraphs[0]
            # выравниваем посередине
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # cells[i].alignment = WD_TABLE_ALIGNMENT.CENTER
            if item == None:
                # цвет текста по свойствам прогона ячейки
                rc = p.runs[0]
                rc.font.color.rgb = RGBColor(255, 0, 0)
                rc.font.bold = True


def AlignParagraph(table, ColList, StartRow, align):
    '''Выравниваем по левому краю ячейку в колонке'''
    for column in ColList:
        for i, cells in enumerate(table.column_cells(column)):
            if i >= StartRow:
                cells.paragraphs[0].alignment = eval(f'WD_ALIGN_PARAGRAPH.{align}')



'''Площадные объекты'''
def dataTable_31(response):
    '''Таблица 3 - Идентификация зданий и сооружений площадочных и линейных объектов'''
    '''Выбираем нужные данные собираем их в список "data"'''
    global response_dop
    data = []
    temp = ''
    DescriptionList = []
    for i in response:
        # КО
        Description = i['KoItem']['Description']
        DescriptionList.append(Description)
        if Description != temp:
            data.append([Description] + [None] * 8)
            temp = Description

        xxx = [
            # № П/П
            i['GenplanNumber'],
            # ЗДАНИЕ/ СООРУЖЕНИЕ
            i['Title'],
            # НАЗНАЧЕНИЕ
            i['Appointment'],
            # ПРИНАДЛЕЖНОСТЬ К ОБЪЕКТАМ ТРАНСПОРТНОЙ ИНФРАСТРУКТУРЫ И К 
            'Да' if i['BelongingTransportFacilities'] == True else 'Нет',
            # ВОЗМОЖНОСТЬ ОПАСНЫХ ПРИРОДНЫХ ПРОЦЕССОВ И ЯВЛЕНИЙ И ТЕХНОГЕННЫХ ВОЗДЕЙСТВИЙ НА ТЕРРИТОРИИ
            # 'HELP!!!',
            response_dop['PossiblityDangerousNaturalProcesses'],
            # ПРИНАДЛЕЖНОСТЬ К ОПАСНЫМ ПРОИЗВОДСТВЕННЫМ ОБЪЕКТАМ
            'Да' if i['BelongingHazardousIndustries'] == True else 'Нет',
            # ПОЖАРНАЯ И ВЗРЫВОПОЖАРНАЯ ОПАСНОСТЬ (Категория здания/сооружения по взрывопожароопасности)
            get_dictionary_item_Title(i['CategoryBuildingFireId']),
            # НАЛИЧИЕ ПОМЕЩЕНИЙ С ПОСТОЯННЫМ ПРЕБЫВАНИЕМ ЛЮДЕЙ 
            'Да' if i['AvailabilityRoomsPeople'] == True else 'Нет',
            # УРОВЕНЬ ОТВЕТСТВЕННОСТИ            '''Неправильно определяется уровень ответственности'''
            get_dictionary_item_Title(i['CategoryLevelResponsibilityId'])
            
        ]
        data.append(xxx)

    DescriptionList = list(set(DescriptionList)) + ['Линейные объекты']
    return data, DescriptionList


'''Линейные объекты'''
def dataTable_32(response):
    '''Таблица 3 - Идентификация зданий и сооружений площадочных и линейных объектов'''
    '''Выбираем нужные данные собираем их в список "data"'''
    global response_dop

    data = [['Линейные объекты'] + [None] * 8]
    
    for i in response:
        xxx = [
            # № П/П
            i['SortIndex'],
            # ЗДАНИЕ/ СООРУЖЕНИЕ
            # i['Title'],
            f"""{i['Title']}\n{i['StartText']}\n{i['FinishText']}""",
            # НАЗНАЧЕНИЕ
            i['PurposeText'],
            # ПРИНАДЛЕЖНОСТЬ К ОБЪЕКТАМ ТРАНСПОРТНОЙ ИНФРАСТРУКТУРЫ И К 
            'Да' if i['IsTransportObject'] == True else 'Нет',
            # ВОЗМОЖНОСТЬ ОПАСНЫХ ПРИРОДНЫХ ПРОЦЕССОВ И ЯВЛЕНИЙ И ТЕХНОГЕННЫХ ВОЗДЕЙСТВИЙ НА ТЕРРИТОРИИ
            response_dop['PossiblityDangerousNaturalProcesses'],
            # ПРИНАДЛЕЖНОСТЬ К ОПАСНЫМ ПРОИЗВОДСТВЕННЫМ ОБЪЕКТАМ
            'Да' if i['IsDangerObject'] == True else 'Нет',
            # ПОЖАРНАЯ И ВЗРЫВОПОЖАРНАЯ ОПАСНОСТЬ (Категория здания/сооружения по взрывопожароопасности)
            get_dictionary_item_Title(i['CategoryBuildingFireId']),
            # НАЛИЧИЕ ПОМЕЩЕНИЙ С ПОСТОЯННЫМ ПРЕБЫВАНИЕМ ЛЮДЕЙ 
            'Да' if i['IsPeopleRoomExists'] == True else 'Нет',
            # УРОВЕНЬ ОТВЕТСТВЕННОСТИ            '''Неправильно определяется уровень ответственности'''
            get_dictionary_item_Title(i['CategoryLevelResponsibilityId'])
        ]
        data.append(xxx)
    # data = sorted(data, key = lambda i: int(i[0]))
    return data


def dataTable_3(response1, response2):
    global doc
    global DescriptionList
    global tableName

    DescriptionList = ['Линейные объекты']

    '''Индексы первой ячейки (строка, стоблец) для вставки данных'''
    if response1 != []:
        data1, DescriptionList = dataTable_31(response1)
    if response2 != []:
        data2 = dataTable_32(response2)

    if response1 != [] and response2 == []:
        data = data1

    if response1 != [] and response2 != []:
        data = [i for g in [data1, data2] for i in g]

    if response1 == [] and response2 != []:
        data = data2

    table = doc.tables[2]

    ExportdataInDoc(table, data)

    '''По колонке ищем первую и последнюю строчку для объединения'''
    column = 4

    '''Объединение ячеек по всей строке таблице с название КО'''
    RowMergeListBase = table.column_cells(0)
    RowMergeList = [i[0] for i in  enumerate(RowMergeListBase) if i[1].text in DescriptionList]

    for row in RowMergeList:
        # объединяем ячейки
        P = MegeCellRows(table, row, 0, 8)
        # делаем шрифт в ячейке жинрый
        P.runs[0].font.bold = True
        # выравниваем по левому краю параграф в ячейке
        P.alignment = WD_ALIGN_PARAGRAPH.LEFT

    '''Объединение ячеек по вертикали с одинаковыми значениями (колонка 5)'''
    celMergeColBase = table.column_cells(column)
    celMergeCol = [ i if t.text not in DescriptionList else None for i, t in  enumerate(celMergeColBase)]
    celMergeCol = celMergeCol[2:]
    celMergeCol.append(None)

    ind_min_max_list = []
    xxx = []

    for v in celMergeCol:
        if v != None:
            xxx.append(v)
        if v == None:
            if xxx != [] and len(xxx) > 1:
                ind_min_max_list.append([min(xxx), max(xxx)])
            xxx = []

    # '''Объединение ячеек в колонках таблице'''
    ColMergeList = [4]
    MegeCellColumns(table, ColMergeList, ind_min_max_list)

    '''Выравниваем по левому краю ячейку в колонке'''
    AlignParagraph(table, ColList = [1, 2], StartRow = 2, align = 'LEFT')


    '''Заменяем текст с номером таблицы в шаблоне '''
    # p_tab_nomer = doc.paragraphs[9]
    # p_tab_nomer.text = f'Таблица {tableNomer}'
    # p_tab_nomer.style = style_right_bold


    NazvanieTab =   f'''Таблица {3}
                    Идентификация зданий и сооружений площадочных и линейных объектов
                    (Федеральный закон № 384 «Технический регламент о безопасности зданий и сооружений»)
                    '''
    tableName.append(' '.join([i.lstrip() for i in NazvanieTab.split('\n')[:3]]))



def dataTable_4(response, tableNomer):
    '''Таблица 4 - Топографическая съемка площадочных объектов'''
    global style_right_bold
    global tableName
    global doc
    
    data = []
    temp = ''
    for i in response:
        # КО
        Description = i['KoItem']['Description']
        if Description != temp:
            data.append(
                [
                Description,
                None,
                response_dop['CharacteristicTerritory'],
                'Сложная конфигурация',
                None,
                response_dop['ApproximateShootingArea'],
                'План 1:500\nПрофиль Мг 1:500,\nМв 1:100,\nМгео 1:100.',
                '0,5',
                response_dop['AdditionalOrSpecialRequirements']
                ]
                )
            temp = Description
            
        # Размеры
        if i['Diameter'] != None:
            try:
                size1 = str(round(float(i['Diameter']) / 1000 + 0.0, 3)).replace('.', ',')
                size2 = size1
            except:
                size1 = size2 = '-'
        else:
            try:
                size1 = str(round(float(i['Length']) / 1000 + 0.0, 3)).replace('.', ',')
            except:
                size1 = '-'
            try:
                size2 = str(round(float(i['Width']) / 1000 + 0.0, 3)).replace('.', ',')
            except:
                size2 = '-'
        
        xxx = [
            # № П/П
            i['GenplanNumber'],
            # ЗДАНИЕ/ СООРУЖЕНИЕ
            i['Title'],
            None,
            # Длина
            size1,
            # Ширина
            size2,
            None,
            None,
            None,
            None
        ]

        data.append(xxx)

    table = doc.tables[3]

    ExportdataInDoc(table, data)

    '''Первая строка с данными нумерация строк начинается с нуля'''
    StartRow = 3
    '''По колонке ищем первую и последнюю строчку для объединения'''
    column = 2
    ind_min_max_list = min_max(table, StartRow, column)
    '''Столбцы, где будем объединять ячейке по типу второго столбца'''
    ColMergeList = [2, 5, 6, 7, 8]

    '''Объединение ячеек в колонках таблице'''
    MegeCellColumns(table, ColMergeList, ind_min_max_list)
    '''Объединение ячеек по всей строке таблице с название КО'''
    RowMergeList = [i[0] for i in ind_min_max_list]
    for row in RowMergeList:
        # объединяем ячейки
        P = MegeCellRows(table, row, 0, 1)
        # делаем шрифт в ячейке жинрый
        P.runs[0].font.bold = True
        # выравниваем по левому краю параграф в ячейке
        P.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # объединяем ячейки (выравнивание параграфа по умолчанию по центре)
        MegeCellRows(table, row, 3, 4)

    '''Выравниваем по левому краю ячейку в колонке'''
    AlignParagraph(table, ColList = [1], StartRow = 3, align = 'LEFT')

    '''Заменяем текст с номером таблицы в шаблоне '''
    p_tab_nomer = doc.paragraphs[13]
    p_tab_nomer.text = f'Таблица {tableNomer}'
    p_tab_nomer.style = style_right_bold

    NazvanieTab =   f'''Таблица {tableNomer}
                    Топографическая съемка площадочных объектов
                    '''
    tableName.append(' '.join([i.lstrip() for i in NazvanieTab.split('\n')[:2]]))


def dataTable_5(response, tableNomer):
    '''Таблица 5 - Топографическая съемка линейных объектов'''
    global doc
    global tableName

    
    data = []
    for i in response:
        xxx = [
            # № по ГП
            i['SortIndex'],
            # НАИМЕНОВАНИЕ ТРАССЫ, ЕЁ НАЧАЛЬНЫЕ И КОНЕЧНЫЕ ПУНКТЫ 
            # i['Title'],
            f"""{i['Title']}\n{i['StartText']}\n{i['FinishText']}""",
            # ПРЕДВАРИТЕЛЬНАЯ ПРОТЯЖЕННОСТЬ ТРАССЫ, КМ
            i['Length'],
            # ШИРИНА ПОЛОСЫ СЪЕМКИ, М
            i['TakeOffWidthText'],
            # МАСШТАБ СЪЕМКИ
            i['TakeOffScaleText'],
            # СЕЧЕНИЕ РЕЛЬЕФА, М
            i['ReliefSize'],
            # МАСШТАБ ПРОДОЛЬНОГО ПРОФИЛЯ
            i['ProfileLongScaleText'],
            # ДОПОЛНИТЕЛЬНЫЕ ИЛИ ОСОБЫЕ ТРЕБОВАНИЯ
            i['AddRequiresText']
        ]
        data.append(xxx)

    table = doc.tables[4]
    '''Отпрвляем данные в таблицу'''
    ExportdataInDoc(table, data)
    '''Выравниваем по левому краю ячейку в колонке'''
    AlignParagraph(table, ColList = [1], StartRow = 2, align = 'LEFT')
    '''Заменяем текст с номером таблицы в шаблоне '''
    p_tab_nomer = doc.paragraphs[16]

    p_tab_nomer.text = f'Таблица {tableNomer}'
    p_tab_nomer.style = style_right_bold

    NazvanieTab =   f'''Таблица {tableNomer}
                    Топографическая съемка линейных объектов
                    '''
    tableName.append(' '.join([i.lstrip() for i in NazvanieTab.split('\n')[:2]]))


def dataTable_6(response, tableNomer):
    '''Таблица 5 - Топографическая съемка линейных объектов'''
    global doc
    global tableName

    data = []
    for i in response:
        xxx = [    
            # № по ГП
            i['SortIndex'],
            # НАИМЕНОВАНИЕ ТРАССЫ 
            # i['Title'],
            f"""{i['Title']}\n{i['StartText']}\n{i['FinishText']}""",
            # ПРОТЯЖЕННОСТЬ ТРАССЫ, КМ
            i['Length'],
            # Столбец 4
            i['BasementDetailsText'],
            # Диаметр
            i['DiameterText'],
            # ДАВЛЕНИЕ
            i['Pressure'],
            # МАТЕРИАЛЬНОЕ ИСПОЛНЕНИЕ
            i['MaterialText'],
            # ОСОБЫЕ УСЛОВИЯ СТРОИТЕЛЬСТВ
            i['SpecialConditionsText']
        ]
        data.append(xxx)

    table = doc.tables[5]
    '''Отпрвляем данные в таблицу'''
    ExportdataInDoc(table, data)
    '''Выравниваем по левому краю ячейку в колонке'''
    AlignParagraph(table, ColList = [1], StartRow = 3, align = 'LEFT')
    '''Заменяем текст с номером таблицы в шаблоне '''
    p_tab_nomer = doc.paragraphs[19]
    p_tab_nomer.text = f'Таблица {tableNomer}'
    p_tab_nomer.style = style_right_bold

    NazvanieTab =   f'''Таблица {tableNomer}
                    Техническая характеристика линейных объектов для инженерно-геологических изысканий
                    '''
    tableName.append(' '.join([i.lstrip() for i in NazvanieTab.split('\n')[:2]]))


def dataTable_7(response, tableNomer):
    global tableName
    global DescriptionList

    data = []
    temp = ''
    for i in response:
        # КО
        Description = i['KoItem']['Description']

        if Description != temp:
            data.append([Description] + [None] * 17)
            temp = Description
        
        col_3 = ''

        if i['Length'] == None and i['Width'] == None and i['Diameter'] != None:
            col_3 = str(round(float(i['Diameter']) / 1000, 3))

        if i['Length'] != None and i['Width'] == None and i['Diameter'] != None:
            col_3 = f"{round(float(i['Length']) / 1000, 3)} x {round(float(i['Diameter']) / 1000, 3)}"

        if i['Length'] == None and i['Width'] != None and i['Diameter'] != None:
            col_3 = f"{round(float(i['Width']) / 1000, 3)} x {round(float(i['Diameter']) / 1000, 3)}"

        if i['Length'] != None and i['Width'] != None and i['Diameter'] == None:
            col_3 = f"{round(float(i['Length']) / 1000, 3)} x {round(float(i['Width']) / 1000, 3)}"

        xxx = [
            # № ЭКСПЛИКАЦИИ ПО СХЕМЕ ГЕНПЛАНА
            i['GenplanNumber'],
            # НАИМЕНОВАНИЕ СООРУЖЕНИЙ
            i['Title'],
            # КОСНТРУКТИВНЫЕ ОСОБЕННОСТИ 
            get_dictionary_item_Title(i['CategoryLandLevelId']),
            # РАЗМЕР В ПЛАНЕ, М
            col_3.replace('.', ','),
            # ОБЩАЯ ВЫСОТА, М
            i['Height'],
            # КОЛИЧЕСТВО ЭТАЖЕЙ
            i['NumberFloorsBuilding'],
            # ОРИЕНТИРОВОЧНАЯ МАССА, Т
            i['Mass'],
            # ТИП (ПЛИТА, ЛЕНТОЧНЫЙ, СВАЙНЫЙ И ДР.)
            get_dictionary_item_Title(i['CategoryFoundationId']),
            # ПРЕДПОЛОГАЕМАЯ ГЛУБИНА ЗАЛОЖЕНИЯ, М
            i['DepthLaying'],
            # СЕЧЕНИЕ СВАЙ, ММ
            i['CrossSectionPiles'],
            # НА ОДНУ СВАЮ (КУСТ СВАЙ), КН (ТС)
            i['LoadPerPile'],
            # НА 1 ПОГОННЫЙ МЕТР ДЛИНЫ ЛЕНТОЧНОГО ФУНДАМЕНТА, КН/М2 (ТС/М2)
            i['LoadPerMeterRibbonFoundation'],
            # ПРЕДПОЛОГАЕМАЯ НА ГРУНТЫ, КН/М2 (ТС/М2)
            i['LoadSoil'],
            # ГЛУБИНА, М
            i['DepthBasement'],
            # НАЗНАЧЕНИЕ
            i['AppointmentBasement'],
            # ДИНАМИЧЕСКИХ НАГРУЗОК
            'Да' if i['AvailabilityDynamicLoads'] == True else 'Нет',
            # МОКРЫХ ТЕХНОЛОГИЧЕСКИХ ПРОЦЕССОВ
            'Да' if i['AvailabilityWetProcesses'] == True else 'Нет',
            # ДОПУСТИМЫЕ ВЕЛИЧИНЫ ДЕФОРМАЦИИ ОСОВАНИЯ, СМ
            '15'
        ]
        data.append(xxx)

    table = doc.tables[6]
    '''Отпрвляем данные в таблицу'''
    ExportdataInDoc(table, data)
    '''Выравниваем по левому краю ячейку в колонке'''
    AlignParagraph(table, ColList = [1], StartRow = 4, align = 'LEFT')

    '''Объединение ячеек по всей строке таблице с название КО'''
    RowMergeListBase = table.column_cells(0)
    RowMergeList = [i[0] for i in  enumerate(RowMergeListBase) if i[1].text in DescriptionList]

    for row in RowMergeList:
        # объединяем ячейки
        P = MegeCellRows(table, row, 0, 17)
        # делаем шрифт в ячейке жинрый
        P.runs[0].font.bold = True
        # выравниваем по центру параграф в ячейке
        P.alignment = WD_ALIGN_PARAGRAPH.CENTER

    '''Заменяем текст с номером таблицы в шаблоне '''
    p_tab_nomer = doc.paragraphs[22]
    p_tab_nomer.text = f'Таблица {tableNomer}'
    p_tab_nomer.style = style_right_bold

    NazvanieTab =   f'''Таблица {tableNomer}
                    Техническая характеристика площадочных объектов для инженерно-геологических изысканий
                    '''
    tableName.append(' '.join([i.lstrip() for i in NazvanieTab.split('\n')[:2]]))


def dataTable_8(response1, response2, tableNomer):
    global tableName

    counter = 0
    data = []
    tempList = []
    if response1 != []:
        for i in response1:
            Description = i['KoItem']['Description']
            if Description not in tempList:
                counter += 1
                tempList.append(Description)
                xxx = [
                    counter,
                    # ИСТОЧНИК ВОЗДЕЙСТВИЯ
                    i['KoItem']['Description'],
                    # РАСПОЛОЖЕНИЕ И ОБЪЕМЫ ИЗЬЯТИЯ ПРИРОДНЫХ РЕСУРСОВ (ЗЕМЕЛЬНЫХ, ВОДНЫХ, ЛЕСНЫХ И Т.Д.)
                    'Земельные участки в пределах отвода на период строительства и эксплуатацию',
                    # ШИРИНА ЗОНЫ ВОЗДЕЙСТВИЯ, М
                    response_dop['KitBuildImpactZoneWidth'],
                    # ГЛУБИНА ВОЗДЕЙСТВИЯ, М
                    response_dop['KitBuildImpactZoneDeep'],
                    # СОСТАВ ЗАГРЯЗНЯЮЩИХ ВЕЩЕСТВ ИЛИ ВИД ВОЗДЕЙСТВИЯ
                    response_dop['KitBuildCompositionPollutants'],
                    # ИНТЕНСИВНОСТЬ И ДЛИТЕЛЬНОСТЬ ВОЗДЕЙСТВИЯ
                    response_dop['KitBuildIntensityDurationExposure']
                ]
                data.append(xxx)
    if response2 != []:
        for i in response2:
            counter += 1
            xxx = [
                counter,
                # ИСТОЧНИК ВОЗДЕЙСТВИЯ
                i['Title'],
                # РАСПОЛОЖЕНИЕ И ОБЪЕМЫ ИЗЬЯТИЯ ПРИРОДНЫХ РЕСУРСОВ (ЗЕМЕЛЬНЫХ, ВОДНЫХ, ЛЕСНЫХ И Т.Д.)
                'Земельные участки в пределах отвода на период строительства и эксплуатацию',
                # ШИРИНА ЗОНЫ ВОЗДЕЙСТВИЯ, М
                response_dop['KitLineImpactZoneWidth'],
                # ГЛУБИНА ВОЗДЕЙСТВИЯ, М
                response_dop['KitLineImpactZoneDeep'],
                # СОСТАВ ЗАГРЯЗНЯЮЩИХ ВЕЩЕСТВ ИЛИ ВИД ВОЗДЕЙСТВИЯ
                response_dop['KitLineCompositionPollutants'],
                # ИНТЕНСИВНОСТЬ И ДЛИТЕЛЬНОСТЬ ВОЗДЕЙСТВИЯ
                response_dop['KitLineIntensityDurationExposure']
            ]
            data.append(xxx)

    table = doc.tables[7]
    '''Отпрвляем данные в таблицу'''
    ExportdataInDoc(table, data)
    '''Выравниваем по левому краю ячейку в колонке'''
    AlignParagraph(table, ColList = [1], StartRow = 2, align = 'LEFT')
    '''Заменяем текст с номером таблицы в шаблоне '''
    p_tab_nomer = doc.paragraphs[25]
    p_tab_nomer.text = f'Таблица {tableNomer}'
    p_tab_nomer.style = style_right_bold

    NazvanieTab =   f'''Таблица {tableNomer}
                    Характеристика существующих и проектируемых источников воздействия
                    '''
    tableName.append(' '.join([i.lstrip() for i in NazvanieTab.split('\n')[:2]]))


def dataTable_1():
    global tableName
    table = doc.tables[0]
    data = [[i + 1, v, ''] for i, v in enumerate(tableName)]
    ExportdataInDoc(table, data)
    AlignParagraph(table, ColList = [1], StartRow = 2, align = 'LEFT')


# @decorTime
def GO(Id, NameFaileDoc):

    global dictionary_items, response_dop
    global tableName, style_right_bold
    global doc
    
    '''Коды словарей'''    
    dictionary_items = resp(r"http://tnnc-pir-app/test-kits-buildings-api/dictionaries/dictionary-items")
    '''Дополнительные данные для формирования ТЗИИ'''
    response_dop = resp(f"http://tnnc-pir-app/test-kits-buildings-api/kits/kit-collection/{Id}")

    '''Наименование объекта'''
    NameObject = response_dop['Project']['Description']

    '''Получаем ответ от сервера'''
    response1 = resp(f"http://tnnc-pir-app/test-kits-buildings-api/kits/kit-build-items/{Id}")
    response1 = [i for i in response1 if i['IsUsedTechnicalSpecificationEngineeringSurvey'] == True]
    response1 = sorted(response1, key = lambda i: (i['KoItem']['Description'], i['GenplanNumber']))
    response2 = resp(f'http://tnnc-pir-app/test-kits-buildings-api/kits/kit-line-items/{Id}')
    
    doc = Document('pattern.docx')
    '''Стиль заоголовка таблицы (жирный справа)'''
    style_right_bold = doc.styles['right_bold']
    '''Задаем стиль текста по умолчанию'''
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'

    '''Список названиий таблиц'''
    tableName = [
        'Таблица 1 Перечень Приложений к ТЗ на ИИ',
        f'Таблица 2 Лист согласования к ТЗ на выполнение ИИ по объекту \"{NameObject}\"'
    ]

    '''Заменяем текст с номером таблицы в шаблоне '''
    p_tab_nomer = doc.paragraphs[7]
    p_tab_nomer.text = f'по объекту \"{NameObject}\"'
    p_tab_nomer.style = style_right_bold

    '''Распечатываем параграфы из шаблона'''
    # for i, val in enumerate(doc.paragraphs): print(i, val.text)

    '''таблица 3'''
    P9 = doc.paragraphs[9]
    P10 = doc.paragraphs[10]
    P11 = doc.paragraphs[11]
    T3 = doc.tables[2]
    '''таблица 4'''
    P13 = doc.paragraphs[13]
    P14 = doc.paragraphs[14]
    P15 = doc.paragraphs[15]
    T4 = doc.tables[3]
    '''таблица 5'''
    P16 = doc.paragraphs[16]
    P17 = doc.paragraphs[17]
    P18 = doc.paragraphs[18]
    T5 = doc.tables[4]
    '''таблица 6'''
    P19 = doc.paragraphs[19]
    P20 = doc.paragraphs[20]
    P21 = doc.paragraphs[21]
    T6 = doc.tables[5]
    '''таблица 7'''
    P22 = doc.paragraphs[22]
    P23 = doc.paragraphs[23]
    P24 = doc.paragraphs[24]
    T7 = doc.tables[6]
    '''таблица 8'''
    P25 = doc.paragraphs[25]
    P26 = doc.paragraphs[26]
    P27 = doc.paragraphs[27]
    T8 = doc.tables[7]

    parara = {
        'Tab3' : [P9 ,P10, P11, T3],
        'Tab4' : [P13 ,P14, P15, T4],
        'Tab5' : [P16 ,P17, P18, T5],
        'Tab6' : [P19 ,P20, P21, T6],
        'Tab7' : [P22 ,P23, P24, T7],
        'Tab8' : [P25 ,P26, P27, T8]
    }

    tableNomer = 3
    if response1 != [] or response2 != []:
        dataTable_3(response1, response2)
        tableNomer += 1
    if response1 != []:
        dataTable_4(response1, tableNomer)
        tableNomer += 1
    if response2 != []:
        dataTable_5(response2, tableNomer)
        tableNomer += 1
        dataTable_6(response2, tableNomer)
        tableNomer += 1
    if response1 != []:
        dataTable_7(response1, tableNomer)
        tableNomer += 1
    dataTable_8(response1, response2, tableNomer)
    dataTable_1()

    if response1 == [] and response2 == []:
        for i in parara['Tab3']: delete_paragraph(i)
    if response1 == []:
        for i in parara['Tab4']: delete_paragraph(i)
        # for i in parara['Tab7']: delete_paragraph(i)
    if response2 == []:
        for i in parara['Tab5']: delete_paragraph(i)
        for i in parara['Tab6']: delete_paragraph(i)

    '''Сохранить как'''
    doc.save(os.getcwd() + f"\\reports\{NameFaileDoc}")



if __name__ == "__main__":
   
    # Id, NameFaileDoc = '2341', 'test_result.docx'
    Id, NameFaileDoc = '2306', 'test_result.docx'
    sys.exit(GO(Id, NameFaileDoc))

